import docx
doc = docx.Document('/home/paarth/development/campus_flow/PRD Document.docx')
for p in doc.paragraphs:
    print(p.text)
for table in doc.tables:
    for row in table.rows:
        print([cell.text for cell in row.cells])
