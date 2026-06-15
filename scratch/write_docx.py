import docx

doc = docx.Document('/home/paarth/development/campus_flow/PRD Document.docx')

replacements = {
    '[What specific problem are you solving? State it in 2-3 sentences. Quantify with data if possible.]': 'Student life is overwhelmingly fragmented. Critical academic, administrative, and personal updates—such as class schedules, assignments, and exam deadlines—are scattered across WhatsApp, emails, and various university portals. This chaos leads to missed deadlines, increased anxiety, and significant wasted time.',
    
    '[Who is affected? How large is the problem? What is the cost of inaction?]': 'Millions of college students globally suffer from information overload and poor executive functioning due to this fragmentation. The cost of inaction is high: missed opportunities, lower academic performance, and severe mental health strain. Solving this streamlines education for Gen-Z.',
    
    '[How does this problem connect to the hackathon theme? What is your unique angle?]': 'Our solution perfectly aligns with the AI for Campus, Community & Everyday Life theme. We are building CampusFlow—an AI operating system that aggregates and synthesizes chaotic data streams into a single, proactive, unified campus assistant.',
    
    '[What has not been tried before? What is the insight that makes your approach different from existing solutions?]': 'Unlike traditional calendar apps or generic university portals that require manual data entry, CampusFlow uses AI to automatically ingest unstructured data (like timetable images and conversational WhatsApp messages), extract deadlines and tasks, and proactively build a personalized schedule and morning briefing.',
    
    '[Who is your user? Describe the persona in 2-3 lines. What do they need?]': 'College and university students who are overwhelmed by juggling academics, extracurriculars, and personal life, and who need a centralized, intelligent system to manage their daily cognitive load.',
    
    '[Describe your solution in plain language. What does it do? Key features (3-5 max):]': 'CampusFlow acts as a smart, unified campus assistant. It centralizes all notifications and schedules, automatically tracks deadlines, and provides proactive daily briefings.',
    
    '[Feature 1]': '1. AI Morning Briefing: A generated daily digest that summarizes urgent tasks, upcoming classes, and missed updates.',
    '[Feature 2]': '2. Automated Timetable & Task Extraction: Uses OCR and LLMs to extract schedules from images and pull deadlines directly from pasted WhatsApp/SMS messages.',
    '[Feature 3]': '3. Smart Voice & Text Task Board: Allows users to log tasks using natural language or voice notes, which the AI automatically categorizes and schedules.',
    
    '[[Insert a simple diagram or 3-step flow showing how a user interacts with your solution]]': '[User opens app] -> [AI Morning Briefing summarizes the day] -> [User uploads timetable/messages] -> [AI automatically structures tasks & deadlines] -> [User tracks progress via Smart Task Board]',
    
    '[[Insert 2-3 screenshots of your working product]]': '[Insert screenshots of CampusFlow UI here]',
    
    '[[Insert system architecture diagram]]': '[Insert CampusFlow Architecture Diagram here]',
    
    '[What core algorithms power your solution? Describe complexity and why you chose this approach.]': 'Our core algorithm relies on Multimodal LLM Extraction pipelines. We use cosine similarity and semantic routing to classify incoming unstructured text (e.g., a WhatsApp message) and extract temporal entities (deadlines) with O(1) latency using Gemini\'s optimized endpoints.',
    
    '[How does this handle 100x-1000x growth? Describe horizontal scaling, caching, or distribution approach.]': 'We utilize a stateless FastAPI backend which can be horizontally scaled via Kubernetes. Heavy AI inference tasks are offloaded to Gemini, meaning our server only handles routing and lightweight data transformations. Image assets are stored in S3-compatible buckets with CDN caching to ensure rapid load times even with 100x user growth.',
    
    '[What does this become in 1-3 years? What is the big-picture vision?]': 'In 1-3 years, CampusFlow will become the definitive "AI OS" for students globally, integrating directly with university LMS systems (like Canvas or Moodle) and serving as a proactive academic advisor.',
    
    '[Which other industries, markets, or user segments can this expand to? Describe the path.]': 'Beyond college students, this technology can expand to K-12 education, corporate onboarding (managing training schedules), and freelance project management, where aggregating unstructured communication into structured tasks is critical.',
    
    '[Quantify: How many users impacted? What cost savings, efficiency gains, or revenue potential at scale?]': 'Impacts millions of students by saving an estimated 5-10 hours a week on scheduling and administrative tasks, significantly reducing academic burnout and improving overall grade outcomes.'
}

# Replace in paragraphs
for p in doc.paragraphs:
    for key, val in replacements.items():
        if key in p.text:
            p.text = p.text.replace(key, val)

# Handle tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            # General text replacements
            for key, val in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, val)
                    
            # Specific table cell overrides
            if cell.text == '[Your Team Name]': cell.text = 'Team CampusFlow'
            if cell.text == '[Theme Name]': cell.text = 'AI for Campus, Community & Everyday Life'
            if cell.text == '[Submission Date]': cell.text = '15 June 2026'
            
            # Members (leave 2 rows)
            if cell.text == '[Member 1]': cell.text = '[Member 1 Name]'
            if cell.text == '[Member 2]': cell.text = '[Member 2 Name]'
            if cell.text == '[Member 3]': cell.text = ''
            if cell.text == '[Member 4]': cell.text = ''
            if cell.text == '[e.g., Backend Dev]': cell.text = '[Role]'
            if cell.text == '[e.g., Frontend Dev]': cell.text = '[Role]'
            if cell.text == '[e.g., ML Engineer]': cell.text = ''
            if cell.text == '[e.g., Designer/DevOps]': cell.text = ''
            if cell.text == '[College]': cell.text = '[College]'
            if cell.text == '[Email]': cell.text = '[Email]'
            
            # Tech stack
            if cell.text == '[e.g., React]': cell.text = 'Flutter (Dart)'
            if cell.text == '[e.g., FastAPI]': cell.text = 'Python FastAPI'
            if cell.text == '[e.g., DynamoDB, Bedrock]': cell.text = 'Google Gemini 1.5 Flash'
            if cell.text == '[e.g., Lambda, S3]': cell.text = 'AWS EC2 & S3'
            
            if 'Frontend' in cell.text and cell.text != 'Frontend': pass
            if cell.text == '[Reason]' and 'Flutter' in row.cells[1].text: cell.text = 'Cross-platform mobile & web deployment'
            if cell.text == '[Reason]' and 'FastAPI' in row.cells[1].text: cell.text = 'High-performance, asynchronous routing'
            if cell.text == '[Reason]' and 'Gemini' in row.cells[1].text: cell.text = 'Rapid multimodal processing (OCR, extraction)'
            if cell.text == '[Reason]' and 'EC2' in row.cells[1].text: cell.text = 'Scalable cloud storage & hosting'
            
            # Roadmap
            if cell.text == '[Near-term]': cell.text = 'Smart notes auto-tagging'
            if cell.text == '[Mid-term]': cell.text = 'Pre-exam readiness checklist & Study group formation'
            if cell.text == '[Long-term]': cell.text = 'Syllabus upload & progress tracker, monthly expense tracker, multi-lingual support'
            
            if cell.text == '[Users/metrics]':
                if '0-3 mo' in row.cells[0].text: cell.text = 'Automates personal note organization'
                if '3-6 mo' in row.cells[0].text: cell.text = 'Boosts exam readiness & peer collaboration'
                if '6-12 mo' in row.cells[0].text: cell.text = 'Full ecosystem integration for diverse users'

doc.save('/home/paarth/development/campus_flow/PRD Document.docx')
